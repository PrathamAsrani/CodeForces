from docx import Document

# Create a new Document
doc = Document()

# Adding Header
doc.add_heading('Pratham Ashok Asrani', 0)
doc.add_paragraph('asranipa@rknec.edu | +91 91454 95032 | LinkedIn Profile')

# Adding Education Section
doc.add_heading('Education', level=1)

edu1 = doc.add_paragraph()
edu1.add_run('Shri Ramdeobaba College of Engineering and Management, Nagpur').bold = True
edu1.add_run('\nBachelors of Technology in Computer Science | CGPA: 9.06 | 2022-2025')

edu2 = doc.add_paragraph()
edu2.add_run('Government Polytechnic Sadar, Nagpur').bold = True
edu2.add_run('\nDiploma in Mechanical Engineering | Aggregate: 95.35% | 2018-2021')

# Adding Skills Section
doc.add_heading('Skills', level=1)

skills = [
    "Programming Languages: C++, Java, Python, Golang",
    "Web Development: React, Express, Mongoose, MongoDB, SQL Server, MySQL",
    "Concepts: OOPS, OS, DBMS, CN",
    "Tools: Git, GitHub, Linux"
]

for skill in skills:
    doc.add_paragraph(skill, style='List Bullet')

# Adding Work Experience Section
doc.add_heading('Work Experience', level=1)

work_exp = doc.add_paragraph()
work_exp.add_run('IEEE Bombay Section, Mumbai').bold = True
work_exp.add_run('\nFull Stack Developer Intern | 2022-2023')
work_exp.add_run('\nProject: Time Tool (Astrology Platform)')
work_exp.add_run('\nTools Used: HTML, CSS, Bootstrap-4, Figma, JavaScript, Firebase')
work_exp.add_run('\nRole: Front-End Developer and Logic Implementation')
work_exp.add_run('\nResponsibilities:')

responsibilities = [
    "Developed an intuitive and visually appealing user interface.",
    "Implemented JavaScript for zodiac sign, leap year, and day count calculations.",
    "Employed Firebase for real-time data management, ensuring responsiveness and cross-device usability."
]

for responsibility in responsibilities:
    doc.add_paragraph(responsibility, style='List Bullet')

# Adding Projects Section
doc.add_heading('Projects', level=1)

project1 = doc.add_paragraph()
project1.add_run('Crypto Tracker').bold = True
project1.add_run('\nFull Stack Developer')
project1.add_run('\nTools Used: React, Ant-Design, Redux, Rapid API')
project1.add_run('\nDescription:')
project1_desc = [
    "Developed a platform for comprehensive cryptocurrency analysis with real-time data.",
    "Implemented interactive line graphs, detailed crypto analysis, and an intuitive search feature."
]
for desc in project1_desc:
    doc.add_paragraph(desc, style='List Bullet')

project2 = doc.add_paragraph()
project2.add_run('Shop Now').bold = True
project2.add_run('\nFull Stack Developer')
project2.add_run('\nTools Used: React, Node.js, Express, MongoDB')
project2.add_run('\nDescription:')
project2_desc = [
    "Developed a full-stack e-commerce website with user and admin roles.",
    "Implemented user authentication, product management, real-time order tracking, and a secure payment gateway."
]
for desc in project2_desc:
    doc.add_paragraph(desc, style='List Bullet')

# Adding Achievements Section
doc.add_heading('Achievements', level=1)

achievements = [
    "Second Rank: Geeks for Geeks",
    "Leetcode Rating: max(1716)",
    "Coding Ninjas Rating: max(2301)",
    "Open Source Contribution: SSOC opensource contributor",
    "Solved over 1000+ problems on Geeks for Geeks and Leetcode"
]
for achievement in achievements:
    doc.add_paragraph(achievement, style='List Bullet')

# Adding Courses Section
doc.add_heading('Courses', level=1)

courses = [
    "Cloud Computing on Azure: Microsoft",
    "Networking Essentials: Cisco",
    "Python Essentials: Cisco",
    "Cloud Services on AWS: Amazon"
]
for course in courses:
    doc.add_paragraph(course, style='List Bullet')

# Save the document to the current working directory
doc.save("Pratham_Ashok_Asrani_Resume.docx")
